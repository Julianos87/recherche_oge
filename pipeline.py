# -*- coding: utf-8 -*-
"""
Pipeline de recherche et d'enrichissement de cabinets de Géomètres-Experts,
puis pré-remplissage des variables pour la rédaction de lettres de motivation.

Sources (toutes publiques, gratuites, sans clé) :
  1. SOURCING  : annuaire officiel de l'Ordre des Géomètres-Experts (OGE).
                 La page /trouver-un-geometre-expert/ embarque la liste complète
                 (~2150 cabinets) en JSON dans l'attribut data-cabinets de #maps.
  2. ENRICH    : API gouvernementale "Recherche d'entreprises"
                 (recherche-entreprises.api.gouv.fr) -> SIRET, dirigeants,
                 tranche d'effectif, date de création, CA, résultat net.
                 Filtre NAF 71.12A = "Activités des géomètres".
  3. SCRAPING  : page-fiche de chaque cabinet sur le site OGE -> spécialités (best effort).
  4. GENERATION: injection des variables dans prompt_lettre.txt -> API Anthropic (optionnel).

Livrables :
  - sortie/cabinets_enrichis.xlsx  (une ligne par cabinet, variables remplies)
  - sortie/lettres/<cabinet>.docx  (si génération activée)

Usage :
  python pipeline.py --departements 87
  python pipeline.py --departements 87 19 23 --lettres
  python pipeline.py --region nouvelle-aquitaine
  python pipeline.py --tout                 (France entière : long !)
"""

import argparse
import html
import json
import os
import re
import sys
import time
import unicodedata
import urllib.parse
import urllib.request
from difflib import SequenceMatcher
from pathlib import Path

import openpyxl
from openpyxl.styles import Font, Alignment
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --------------------------------------------------------------------------- #
# Constantes
# --------------------------------------------------------------------------- #
RACINE = Path(__file__).resolve().parent
SORTIE = RACINE / "sortie"
LETTRES = SORTIE / "lettres"
CACHE_HTML = SORTIE / "_oge_annuaire.html"

URL_ANNUAIRE = "https://www.geometre-expert.fr/trouver-un-geometre-expert/"
URL_API_ENTREPRISES = "https://recherche-entreprises.api.gouv.fr/search"
NAF_GEOMETRE = "71.12A"  # Activités des géomètres
UA = "Mozilla/5.0 (recherche-stage-DPLG; usage personnel)"

# Départements par région (codes INSEE), pour le filtre --region
REGIONS = {
    "auvergne-rhone-alpes": ["01", "03", "07", "15", "26", "38", "42", "43", "63", "69", "73", "74"],
    "bourgogne-franche-comte": ["21", "25", "39", "58", "71", "70", "89", "90"],
    "bretagne": ["22", "29", "35", "56"],
    "centre-val-de-loire": ["18", "28", "36", "37", "41", "45"],
    "corse": ["20"],
    "grand-est": ["08", "10", "51", "52", "54", "55", "57", "67", "68", "88"],
    "hauts-de-france": ["02", "59", "60", "62", "80"],
    "ile-de-france": ["75", "77", "78", "91", "92", "93", "94", "95"],
    "normandie": ["14", "27", "50", "61", "76"],
    "nouvelle-aquitaine": ["16", "17", "19", "23", "24", "33", "40", "47", "64", "79", "86", "87"],
    "occitanie": ["09", "11", "12", "30", "31", "32", "34", "46", "48", "65", "66", "81", "82"],
    "pays-de-la-loire": ["44", "49", "53", "72", "85"],
    "provence-alpes-cote-d-azur": ["04", "05", "06", "13", "83", "84"],
}

# Mots du nom de cabinet à ignorer pour deviner le patronyme à interroger
STOP = {"sarl", "selarl", "scp", "sas", "sasu", "eurl", "sa", "et", "associes",
        "associés", "cabinet", "geometre", "geometres", "expert", "experts",
        "de", "du", "des", "le", "la", "les", "bs", "pe", "groupe", "atelier",
        "bureau", "société", "societe", "selas", "scop", "mesures"}


# --------------------------------------------------------------------------- #
# Utilitaires
# --------------------------------------------------------------------------- #
def http_get(url, params=None, timeout=25):
    import ssl
    if params:
        url = url + "?" + urllib.parse.urlencode(params)
    # User-Agent imitant un vrai navigateur pour éviter d'être bloqué
    ua = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    req = urllib.request.Request(url, headers={"User-Agent": ua})
    
    # Contexte SSL permissif pour les sites de cabinets mal configurés
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    
    with urllib.request.urlopen(req, timeout=timeout, context=ctx) as r:
        raw = r.read()
    return raw


def api_entreprises(params, retries=4):
    """Appel résilient à l'API entreprises (retry + backoff sur erreurs/429).

    L'API limite à ~7 requêtes/seconde par IP : garde une pause >= 0.2 s entre
    deux appels pour ne pas se faire bannir temporairement.
    """
    import urllib.error
    delai = 1.0
    for tentative in range(retries):
        try:
            return json.loads(http_get(URL_API_ENTREPRISES, params, timeout=20)
                              .decode("utf-8"))
        except urllib.error.HTTPError as e:
            if e.code == 429 and tentative < retries - 1:   # trop de requêtes
                time.sleep(delai)
                delai *= 2
                continue
            raise
        except Exception:
            if tentative < retries - 1:
                time.sleep(delai)
                delai *= 2
                continue
            raise
    return {"results": []}


def sans_accent(s):
    return "".join(c for c in unicodedata.normalize("NFD", s or "")
                   if unicodedata.category(c) != "Mn")


def similarite(a, b):
    return SequenceMatcher(None, sans_accent(a).lower(),
                           sans_accent(b).lower()).ratio()


# --------------------------------------------------------------------------- #
# 1. SOURCING
# --------------------------------------------------------------------------- #
def charger_cabinets(force=False):
    """Renvoie la liste complète des cabinets depuis l'annuaire OGE."""
    if CACHE_HTML.exists() and not force:
        donnees = CACHE_HTML.read_bytes()
    else:
        print("Téléchargement de l'annuaire OGE...")
        donnees = http_get(URL_ANNUAIRE)
        CACHE_HTML.write_bytes(donnees)
    # La page mêle UTF-8 / Windows-1252 dans les adresses : on décode en
    # tolérant pour ne pas perdre de cabinet.
    texte = donnees.decode("utf-8", errors="replace")
    m = re.search(r"data-cabinets=(['\"])(\[.*?\])\1", texte, re.S)
    if not m:
        raise RuntimeError("Bloc data-cabinets introuvable (le site a changé ?).")
    cabinets = json.loads(html.unescape(m.group(2)))
    print(f"{len(cabinets)} cabinets trouvés dans l'annuaire.")
    return cabinets


def filtrer(cabinets, departements=None, sieges_seuls=True):
    res = []
    for c in cabinets:
        if sieges_seuls and c.get("isbureau"):
            continue  # on écarte les bureaux secondaires
            
        cp = str(c.get("zipcode", ""))
        # Exclusion des DOM-TOM (97, 98) et étranger (99) pour la France Hexagonale
        if departements is None and (cp.startswith("97") or cp.startswith("98") or cp.startswith("99")):
            continue
            
        if departements:
            if not any(cp.startswith(d) for d in departements):
                continue
        res.append(c)
    return res


# --------------------------------------------------------------------------- #
# 2. ENRICHISSEMENT (API Recherche d'entreprises)
# --------------------------------------------------------------------------- #
def patronyme_recherche(label):
    """Devine le terme le plus discriminant du nom (souvent le patronyme)."""
    mots = re.findall(r"[A-Za-zÀ-ÿ]+", label)
    candidats = [m for m in mots if sans_accent(m).lower() not in STOP and len(m) > 2]
    return candidats[0] if candidats else (mots[0] if mots else label)


def enrichir(cabinet):
    """Interroge l'API entreprises et renvoie un dict d'infos (best effort)."""
    cp = str(cabinet.get("zipcode", "")).strip()
    terme = patronyme_recherche(cabinet["label"])
    info = {"siren": "", "naf": "", "dirigeants": "", "effectif": "",
            "creation": "", "ca": "", "resultat": "", "infos_financieres": ""}
    try:
        params = {"q": terme, "per_page": 10}
        if len(cp) == 5:
            params["code_postal"] = cp
        rep = api_entreprises(params)
    except Exception as e:
        info["infos_financieres"] = "Non disponible"
        return info

    # On ne garde que les géomètres, puis on prend le meilleur match de nom.
    candidats = [e for e in rep.get("results", [])
                 if (e.get("activite_principale") or "").startswith("71.12")]
    if not candidats:
        info["infos_financieres"] = "Non disponible"
        return info
    meilleur = max(candidats, key=lambda e: similarite(e.get("nom_complet", ""),
                                                        cabinet["label"]))

    info["siren"] = meilleur.get("siren", "")
    info["naf"] = meilleur.get("activite_principale", "")
    info["creation"] = (meilleur.get("date_creation") or "")[:4]

    dirs = []
    for d in meilleur.get("dirigeants", []):
        nom = " ".join(x for x in [(d.get("prenoms") or "").title(),
                                   (d.get("nom") or "").title()] if x).strip()
        if d.get("denomination"):
            nom = d["denomination"]
        if nom and nom not in dirs:
            dirs.append(nom)
    info["dirigeants"] = " ; ".join(dirs[:3])

    eff = meilleur.get("tranche_effectif_salarie")
    info["effectif"] = libelle_effectif(eff)

    fin = meilleur.get("finances") or {}
    if fin:
        annee = sorted(fin.keys())[-1]
        info["ca"] = fin[annee].get("ca")
        info["resultat"] = fin[annee].get("resultat_net")
        info["_annee_fin"] = annee

    info["infos_financieres"] = phrase_financiere(info)
    return info


def libelle_effectif(code):
    table = {
        "NN": "", "00": "0 salarié", "01": "1 à 2 salariés", "02": "3 à 5 salariés",
        "03": "6 à 9 salariés", "11": "10 à 19 salariés", "12": "20 à 49 salariés",
        "21": "50 à 99 salariés", "22": "100 à 199 salariés",
    }
    return table.get(code, "")


def phrase_financiere(info):
    """Compile une phrase propre pour la variable {infos_financieres}."""
    bouts = []
    if info.get("creation"):
        bouts.append(f"créé en {info['creation']}")
    if info.get("effectif"):
        bouts.append(info["effectif"])
    ca = info.get("ca")
    if ca:
        annee = info.get("_annee_fin", "")
        bouts.append(f"CA de {int(ca):,} €".replace(",", " ") +
                     (f" ({annee})" if annee else ""))
    return ", ".join(bouts).capitalize() if bouts else "Non disponible"


# --------------------------------------------------------------------------- #
# 3. SPÉCIALITÉS : découverte du site propre du cabinet + scraping ciblé
# --------------------------------------------------------------------------- #
# Le site OGE affiche un gabarit listant TOUTES les activités -> inutilisable.
# On cherche donc le site PROPRE du cabinet, déduit du domaine de son email
# (les fournisseurs génériques ci-dessous ne sont pas des sites de cabinet).
FOURNISSEURS_GENERIQUES = {
    "wanadoo.fr", "orange.fr", "gmail.com", "free.fr", "sfr.fr", "laposte.net",
    "hotmail.fr", "hotmail.com", "outlook.fr", "outlook.com", "yahoo.fr",
    "yahoo.com", "geometre-expert.fr", "live.fr", "neuf.fr", "bbox.fr",
    "aliceadsl.fr", "numericable.fr", "geomesure.fr",
}

MOTS_SPECIALITES = [
    "bornage", "copropriété", "division", "lotissement", "aménagement",
    "urbanisme", "topographie", "3D", "scan", "lidar", "lazer", "laser",
    "foncier", "vrd", "maîtrise d'œuvre", "drone", "photogrammétrie",
    "diagnostic", "implantation", "récolement", "cadastre", "servitude",
    "modélisation", "BIM", "expertise", "géodésie", "nivellement",
]

# Pages internes fréquentes où sont décrites les prestations
PAGES_CANDIDATES = ["", "/nos-prestations", "/prestations", "/nos-services",
                    "/services", "/competences", "/nos-competences",
                    "/savoir-faire", "/metiers", "/activites"]


def decouvrir_site(cabinet):
    """Déduit l'URL du site propre du cabinet à partir du domaine email."""
    email = (cabinet.get("email") or "").strip().lower()
    if "@" not in email:
        return ""
    domaine = email.split("@", 1)[1]
    if domaine in FOURNISSEURS_GENERIQUES:
        return ""
    return "https://" + domaine


def scraper_specialites(cabinet):
    """Scrape le site propre du cabinet et renvoie ses spécialités déclarées."""
    site = decouvrir_site(cabinet)
    if not site:
        return "Non disponible"
    texte_total = ""
    for chemin in PAGES_CANDIDATES[:5]:  # homepage + 4 pages probables
        try:
            texte_total += " " + http_get(site + chemin, timeout=12).decode(
                "utf-8", errors="replace")
        except Exception:
            continue
    if not texte_total.strip():
        return "Non disponible"
    bas = sans_accent(texte_total).lower()
    trouves = [m for m in MOTS_SPECIALITES
               if sans_accent(m).lower() in bas]
    return ", ".join(trouves) if trouves else "Non disponible"


# --------------------------------------------------------------------------- #
# 3 bis. RECRUTEMENT : API France Travail (Offres d'emploi)
# --------------------------------------------------------------------------- #
# Gratuite mais nécessite d'enregistrer une application sur https://francetravail.io
# (API "Offres d'emploi v2"), puis d'exporter :
#     FT_CLIENT_ID, FT_CLIENT_SECRET
FT_TOKEN_URL = ("https://entreprise.francetravail.fr/connexion/oauth2/"
                "access_token?realm=%2Fpartenaire")
FT_SEARCH_URL = "https://api.francetravail.io/partenaire/offresdemploi/v2/offres/search"
FT_SCOPE = "api_offresdemploiv2 o2dsoffre"


def ft_token():
    """Récupère un jeton OAuth2 (client_credentials) ou None si non configuré."""
    cid = os.environ.get("FT_CLIENT_ID")
    secret = os.environ.get("FT_CLIENT_SECRET")
    if not cid or not secret:
        return None
    data = urllib.parse.urlencode({
        "grant_type": "client_credentials",
        "client_id": cid,
        "client_secret": secret,
        "scope": FT_SCOPE,
    }).encode()
    req = urllib.request.Request(FT_TOKEN_URL, data=data, headers={
        "Content-Type": "application/x-www-form-urlencoded"})
    try:
        rep = json.loads(urllib.request.urlopen(req, timeout=20).read())
        return rep.get("access_token")
    except Exception as e:
        print("  [France Travail] échec token:", e)
        return None


def offres_recrutement(cabinet, token):
    """Renvoie une phrase {infos_recrutement} d'après les offres détectées."""
    if not token:
        return "Non disponible"
    cp = str(cabinet.get("zipcode", ""))
    dep = cp[:2] if len(cp) >= 2 else ""
    terme = patronyme_recherche(cabinet["label"])
    params = {"motsCles": terme, "range": "0-9"}
    if dep:
        params["departement"] = dep
    url = FT_SEARCH_URL + "?" + urllib.parse.urlencode(params)
    req = urllib.request.Request(url, headers={
        "Authorization": "Bearer " + token, "Accept": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=20) as r:
            if r.status == 204:           # aucun résultat
                return "Non disponible"
            rep = json.loads(r.read())
    except Exception as e:
        return "Non disponible"

    offres = rep.get("resultats", [])
    # On ne garde que les offres dont l'entreprise correspond au cabinet.
    pertinentes = []
    for o in offres:
        nom_ent = (o.get("entreprise", {}) or {}).get("nom", "") or ""
        if similarite(nom_ent, cabinet["label"]) > 0.45 or \
           sans_accent(terme).lower() in sans_accent(nom_ent).lower():
            pertinentes.append(o.get("intitule", "").strip())
    if not pertinentes:
        return "Non disponible"
    intitules = " ; ".join(dict.fromkeys(p for p in pertinentes if p))[:200]
    n = len(pertinentes)
    return f"Recrute actuellement ({n} offre{'s' if n > 1 else ''} : {intitules})"


# --------------------------------------------------------------------------- #
# Assemblage des variables d'un cabinet
# --------------------------------------------------------------------------- #
def construire_variables(cabinet, ent, specialites, recrutement="Non disponible"):
    nom = cabinet["label"]
    adresse = ", ".join(x for x in [cabinet.get("fullAddress", "").strip(),
                                    cabinet.get("fullCity", "").strip()] if x)
    # Le dirigeant nominatif vient de l'API ; sinon on n'invente rien.
    dirigeant = ent.get("dirigeants") or "Non disponible"
    return {
        "nom_cabinet": nom,
        "nom_dirigeant": dirigeant,
        "adresse_cabinet": adresse or "Non disponible",
        "infos_financieres": ent.get("infos_financieres", "Non disponible"),
        "specialites_scrapees": specialites,
        "infos_recrutement": recrutement,
        "autres_infos": "Non disponible",
        # champs techniques utiles dans le tableur
        "_email": cabinet.get("email", ""),
        "_tel": cabinet.get("phone", ""),
        "_lien": cabinet.get("link", ""),
        "_siren": ent.get("siren", ""),
    }


# --------------------------------------------------------------------------- #
# 4. GÉNÉRATION DE LETTRE (optionnelle : nécessite ANTHROPIC_API_KEY)
# --------------------------------------------------------------------------- #
def generer_lettre(variables, profil, gabarit):
    try:
        import anthropic
    except ImportError:
        return None
    cle = os.environ.get("ANTHROPIC_API_KEY")
    if not cle:
        return None
    prompt = gabarit.format(profil_candidat=profil, **{
        k: variables[k] for k in
        ["nom_cabinet", "nom_dirigeant", "adresse_cabinet", "infos_financieres",
         "specialites_scrapees", "infos_recrutement", "autres_infos"]})
    client = anthropic.Anthropic(api_key=cle)
    msg = client.messages.create(
        model="claude-opus-4-8",
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}],
    )
    return "".join(b.text for b in msg.content if getattr(b, "type", "") == "text")


def ecrire_docx(texte, chemin):
    doc = Document()
    
    # Marges strictes
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)
        
    # On extrait uniquement les lignes avec du texte pour ignorer les sauts aléatoires de l'IA
    lignes_brutes = texte.split("\n")
    lignes = [l.strip() for l in lignes_brutes if l.strip()]
    
    # Trouver l'indice de l'objet
    idx_objet = -1
    for i, l in enumerate(lignes):
        if l.lower().startswith("objet"):
            idx_objet = i
            break
            
    # Fallback au cas où l'IA oublie le mot "Objet"
    if idx_objet == -1:
        idx_objet = 8
        
    # On s'assure que la lettre finit par Julian Brouet et on coupe ce qu'il y a après
    for i in range(len(lignes)-1, -1, -1):
        if "julian brouet" in lignes[i].lower():
            lignes = lignes[:i+1]
            break

    idx_signature = len(lignes) - 1
    
    for i, l in enumerate(lignes):
        if i == idx_objet:
            # Forcer exactement 2 lignes vides avant l'objet
            para = doc.add_paragraph("")
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            para = doc.add_paragraph("")
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0

        para = doc.add_paragraph(l)
        
        if i < 5:
            # Bloc 1: Expéditeur
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            
        elif i >= 5 and i < idx_objet:
            # Bloc 2: Destinataire (avant Objet)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            
        elif i == idx_objet:
            # Objet
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            
            # Forcer exactement 1 ligne vide après l'objet
            para_vide = doc.add_paragraph("")
            para_vide.paragraph_format.space_before = Pt(0)
            para_vide.paragraph_format.space_after = Pt(0)
            para_vide.paragraph_format.line_spacing = 1.15
            
        elif i == idx_signature:
            # Signature (Julian Brouet final)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            
        else:
            # Corps du texte
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(8)  # Espacement natif de Word entre les paragraphes
            
    doc.save(chemin)


# --------------------------------------------------------------------------- #
# Export tableur
# --------------------------------------------------------------------------- #
COLONNES = [
    ("nom_cabinet", "Cabinet"), ("nom_dirigeant", "Dirigeant (GE)"),
    ("adresse_cabinet", "Adresse"), ("infos_financieres", "Infos financières"),
    ("specialites_scrapees", "Spécialités"), ("infos_recrutement", "Recrutement"),
    ("_email", "Email"), ("_tel", "Téléphone"), ("_siren", "SIREN"),
    ("_lien", "Fiche OGE"),
]


def exporter_xlsx(lignes, chemin):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Cabinets"
    for j, (_, titre) in enumerate(COLONNES, 1):
        c = ws.cell(1, j, titre)
        c.font = Font(bold=True)
    for i, var in enumerate(lignes, 2):
        for j, (cle, _) in enumerate(COLONNES, 1):
            ws.cell(i, j, var.get(cle, ""))
    # largeurs
    larg = [32, 26, 40, 34, 30, 14, 22, 16, 12, 40]
    for j, w in enumerate(larg, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(j)].width = w
    ws.freeze_panes = "A2"
    wb.save(chemin)


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #
def main():
    ap = argparse.ArgumentParser(description="Recherche de cabinets GE + variables lettres.")
    ap.add_argument("--departements", nargs="+", help="codes département, ex: 87 19 23")
    ap.add_argument("--region", choices=sorted(REGIONS), help="région prédéfinie")
    ap.add_argument("--tout", action="store_true", help="France entière (long)")
    ap.add_argument("--specialites", action="store_true",
                    help="scraper le site propre du cabinet pour ses spécialités")
    ap.add_argument("--recrutement", action="store_true",
                    help="détecter les offres via France Travail (FT_CLIENT_ID/SECRET)")
    ap.add_argument("--lettres", action="store_true",
                    help="générer les lettres .docx (nécessite ANTHROPIC_API_KEY)")
    ap.add_argument("--pause", type=float, default=0.3, help="pause entre appels API (s)")
    ap.add_argument("--max", type=int, default=0, help="limiter le nombre de cabinets (test)")
    args = ap.parse_args()

    SORTIE.mkdir(exist_ok=True)
    LETTRES.mkdir(exist_ok=True)

    departements = None
    if args.region:
        departements = REGIONS[args.region]
    elif args.departements:
        departements = [d.zfill(2) for d in args.departements]
    elif not args.tout:
        ap.error("Précise --departements, --region ou --tout.")

    cabinets = filtrer(charger_cabinets(), departements)
    if args.max:
        cabinets = cabinets[:args.max]
    print(f"{len(cabinets)} cabinets à traiter.")

    profil = (RACINE / "profil_candidat.txt").read_text(encoding="utf-8")
    gabarit = (RACINE / "prompt_lettre.txt").read_text(encoding="utf-8")

    token_ft = None
    if args.recrutement:
        token_ft = ft_token()
        if not token_ft:
            print("  [France Travail] identifiants absents/invalides "
                  "(FT_CLIENT_ID / FT_CLIENT_SECRET) -> recrutement ignoré.")

    lignes = []
    for n, cab in enumerate(cabinets, 1):
        print(f"  [{n}/{len(cabinets)}] {cab['label']}")
        ent = enrichir(cab)
        time.sleep(args.pause)
        spec = scraper_specialites(cab) if args.specialites else "Non disponible"
        recr = offres_recrutement(cab, token_ft) if args.recrutement else "Non disponible"
        var = construire_variables(cab, ent, spec, recr)
        lignes.append(var)

        if args.lettres:
            texte = generer_lettre(var, profil, gabarit)
            if texte:
                slug = re.sub(r"[^a-z0-9]+", "-",
                              sans_accent(cab["label"]).lower()).strip("-")[:60]
                ecrire_docx(texte, LETTRES / f"{slug}.docx")

    chemin_xlsx = SORTIE / "cabinets_enrichis.xlsx"
    exporter_xlsx(lignes, chemin_xlsx)
    print(f"\nTerminé. Tableur : {chemin_xlsx}")
    if args.lettres:
        print(f"Lettres : {LETTRES}")


if __name__ == "__main__":
    main()
